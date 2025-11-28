import streamlit as st
import random
from streamlit_calendar import calendar
import uuid
import time
from datetime import date, timedelta, datetime
import openai
import google.generativeai as genai 
from st_supabase_connection import SupabaseConnection
from supabase import create_client, Client
import supabase
import json
from docx import Document
import io
import os
import base64
from pathlib import Path


# Initialize Supabase Client
url = st.secrets['connections']['SUPABASE_URL']
key = st.secrets['connections']['SUPABASE_KEY']
supabase: Client = create_client(url, key)

# Setting the Background
def background():
    page_element="""
    <style>
    [data-testid="stAppViewContainer"]{
    background-image: url("https://wallpapers.com/images/featured/dark-mountain-gd3o1mx0wxezbewk.jpg");
    background-size: cover;
    }
    [data-testid="stSidebar"] {
    background: rgba(0, 0, 0, 0);
    }
    [data-testid="stSidebar"]> div:first-child{
    background-image: url("https://wallpapers.com/images/featured/dark-mountain-gd3o1mx0wxezbewk.jpg");
    background-size: cover;
    }
    </style>


    """

    st.markdown(page_element, unsafe_allow_html=True)
background()

# Frosted Glass Theme CSS
st.markdown("""
<style>
/* --- Buttons --- */
button[kind="secondary"], button[kind="primary"], div.stButton > button {
    background-color: rgba(255, 255, 255, 0.08) !important; /* semi-transparent */
    color: white !important;
    border: 1px solid rgba(255, 255, 255, 0.3) !important;
    backdrop-filter: blur(6px);
    transition: 0.3s ease-in-out;
    border-radius: 8px;
}

/* Button hover */
div.stButton > button:hover {
    background-color: rgba(255, 255, 255, 0.25) !important;
    border: 1px solid rgba(255, 255, 255, 0.5) !important;
}

/* --- Dialogs / Modals --- */
.stDialog, .stModal, div[data-testid="stModal"], div[data-testid="stDialog"] {
    background-color: rgba(30, 30, 30, 0.4) !important;
    color: white !important;
    backdrop-filter: blur(10px);
    border-radius: 12px;
    border: 1px solid rgba(255, 255, 255, 0.25);
}

/* Inner containers in dialogs */
.stDialog div, .stModal div {
    background-color: transparent !important;
}

/* --- Inputs / Textareas / Selects --- */
input, select, textarea {
    background-color: rgba(255, 255, 255, 0.1) !important;
    color: white !important;
    border: 1px solid rgba(255, 255, 255, 0.3) !important;
    border-radius: 6px;
}

/* --- Popovers (expanded content) --- */
div[data-baseweb="popover"] {
    background-color: rgba(30, 30, 30, 0.4) !important; 
    color: white !important;
    border-radius: 12px !important;
    border: 1px solid rgba(255, 255, 255, 0.25) !important;
    box-shadow: none !important;
    backdrop-filter: blur(10px) !important;
}

/* --- Top Header / Nav Bar --- */
header[data-testid="stHeader"] {
    background-color: rgba(30, 30, 30, 0.4) !important; /* frosted glass */
    backdrop-filter: blur(10px);
    border-bottom: none !important;
    box-shadow: none !important;
}

/* Keep icons and text visible */
header[data-testid="stHeader"] * {
    color: white !important;
}

/* --- Buttons in the top bar --- */
header[data-testid="stHeader"] button {
    background-color: rgba(255, 255, 255, 0.08) !important; /* semi-transparent */
    color: white !important;
    border: 1px solid rgba(255, 255, 255, 0.3) !important;
    backdrop-filter: blur(6px);
    border-radius: 8px;
    padding: 6px 12px;
    font-weight: bold;
    transition: 0.3s ease-in-out;
}

/* Hover effect for buttons */
header[data-testid="stHeader"] button:hover {
    background-color: rgba(255, 255, 255, 0.25) !important;
    border: 1px solid rgba(255, 255, 255, 0.5) !important;
}

/* Optional: make main content match frosted theme */
.main .block-container {
    background-color: rgba(30, 30, 30, 0.3) !important;
    border-radius: 12px;
    padding: 20px;
    backdrop-filter: blur(10px);
}

/* Popover inner elements */
div[data-baseweb="popover"] * {
    background-color: transparent !important;
    color: white !important;
}

/* Optional: Adjust popover arrow to match background */
div[data-baseweb="popover"]::after {
    background-color: rgba(30, 30, 30, 0.4) !important;
}
            


/* Button hover effect */
div.stButton > button:hover {
    background-color: rgba(255, 255, 255, 0.2) !important;
    border: 2px solid rgba(255, 255, 255, 0.7) !important;
}

/* Transparent container with borders and frosted glass */
.task-container {
    background-color: rgba(0, 0, 0, 0) !important; /* fully transparent */
    color: white !important;
    border-radius: 12px;
    border: 1px solid rgba(255, 255, 255, 0.5); /* visible border */
    padding: 15px;
    margin-top: 5px;
    margin-bottom: 10px;
    backdrop-filter: blur(10px); /* frosted effect */
}

/* --- Frosted Expanders --- */
div.streamlit-expanderHeader {
    background: rgba(255, 255, 255, 0.12) !important;
    backdrop-filter: blur(10px) !important;
    border: 1px solid rgba(255, 255, 255, 0.25) !important;
    border-radius: 10px !important;
    color: white !important;
    transition: all 0.3s ease-in-out !important;
}


/* Hover / active glow for expander header */
div.streamlit-expanderHeader:hover {
    background: rgba(255, 255, 255, 0.22) !important;
    border-color: rgba(255, 255, 255, 0.5) !important;
}

/* Remove Streamlit default shadows */
div[data-testid="stExpander"] {
    box-shadow: none !important;
}
            
/* --- Frosted Expanders --- */
.frosted-expander-header {
    background: rgba(255, 255, 255, 0.08);
    backdrop-filter: blur(10px);
    border: 1px solid rgba(255, 255, 255, 0.25);
    border-radius: 10px;
    padding: 12px 16px;
    margin-top: 10px;
    color: white;
    cursor: pointer;
    font-weight: 500;
    transition: all 0.3s ease-in-out;
}
.frosted-expander-header:hover {
    background: rgba(255, 255, 255, 0.18);
    border-color: rgba(255, 255, 255, 0.4);
}

.frosted-expander-content {
    background: rgba(255, 255, 255, 0.06);
    backdrop-filter: blur(10px);
    border: 1px solid rgba(255, 255, 255, 0.15);
    border-top: none;
    border-radius: 0 0 10px 10px;
    padding: 0 15px;
    color: white;
    max-height: 0;
    overflow: hidden;
    transition: max-height 0.4s ease, padding 0.3s ease;
}
.frosted-expander-content.show {
    max-height: 1000px; /* Large enough to show full content */
    padding: 15px;
}
</style>

<script>
function toggleExpander(index) {
    const all = document.querySelectorAll('.frosted-expander-content');
    all.forEach(div => {
        if (div.id === `expander-${index}`) {
            div.classList.toggle('show');
        } else {
            div.classList.remove('show');
        }
    });
}
</script>
""", unsafe_allow_html=True)


# Setting The User Session
def get_user_data(email):
    # Fetch user data from Supabase or initialize if not found
    response = supabase.table("user_data").select("data").eq("email", email).execute()
    data = response.data

    if not data:
        # User doesn't exist yet, initialize
        default_data = {
            "courses_list": [],
            "courses_colors": [],
            "difficulty_ranking": [],
            "tasks": [],
            "complete_tasks": [],
            "written_notes":[ ],
            "uploaded_file":[],
            "ai_history":[],
            "ai_use_task_ordering":[],
            "ai_use_ai_priority":[],
            "ai_document_assistant":[],
            "ai_use_history":[],
            "ai_quiz_length": ["Short"],
            "ai_summary_length": ["Short"],
            "ai_assistant_response_length": ["Medium"],

        }

        supabase.table("user_data").insert({
            "email": email,
            "data": default_data
        }).execute()

        return default_data

    # Ensures that the data coming to the app is a string and what goes to the database is json compatible 

    raw_data = data[0]["data"]
    if isinstance(raw_data, str):
        return json.loads(raw_data)
    return raw_data



# Initialize sectioned task list session state
if "today_tasks" not in st.session_state:
    st.session_state["today_tasks"] = []

if "this_week_tasks" not in st.session_state:
    st.session_state["this_week_tasks"] = []

if "ai_to_do_list_database" not in st.session_state:
    st.session_state["ai_to_do_list_database"] = []


def get_week_bounds():
    """Return today's date, start of week, and end of week (all as date objects)."""
    today = datetime.today().date()
    start_of_week = today - timedelta(days=today.isoweekday() - 1)  # Monday start
    end_of_week = start_of_week + timedelta(days=6)
    return today, start_of_week, end_of_week


user_data = get_user_data(st.session_state['username'])
today, start, end = get_week_bounds()

def parse(d):
        return datetime.fromisoformat(d).date()  # all tasks stored as strings
todays_tasks = [
        task["name"]
        for task in user_data['tasks']
        if parse(task["due_date"]) == today
    ]

this_weeks_tasks = [
    task["name"]
    for task in user_data['tasks']
    if start <= parse(task["due_date"]) <= end and parse(task["due_date"]) != today
]


# Function to Update Supabase Database 

def update_user_data(email, new_data):
    """Update user's data field in Supabase."""
    supabase.table("user_data").update({
        "data": new_data
    }).eq("email", email).execute()


def ai(email):


    # Fetch user data
    user_data = get_user_data(email)

    # Get the data from database to edit within application
    courses = user_data.get("courses_list", [])
    colors = user_data.get("courses_colors", [])
    difficulty_ranking = user_data.get("difficulty_ranking", [])
    tasks = user_data.get("tasks", [])
    completed_tasks = user_data.get("complete_tasks", [])
    written_notes = user_data.get("written_notes", [])
    uploaded_file = user_data.get("uploaded_file", [])
    the_ai_history = user_data.get("ai_history", [])
    ai_quiz_length = user_data.get("ai_quiz_length", [])
    ai_summary_length = user_data.get("ai_summary_length", [])
    ai_assistant_response_length = user_data.get("ai_assistant_response_length", [])


    if "ai_quiz" not in st.session_state:
            st.session_state["ai_quiz"]  = []
    if "ai_data_stale" not in st.session_state:
            st.session_state["ai_data_stale"] = True
    if "ai_data_stale_priority" not in st.session_state:
            st.session_state["ai_data_stale_priority"] = True
    if "ai_priority" not in st.session_state:
        st.session_state['ai_priority'] = "To Be Determined"

    def display_tasks():
        global username
        global today

        # Sort tasks: overdue first, then upcoming in chronological order
        sorted_tasks = sorted(
            user_data["tasks"],
            key=lambda t: (
                datetime.fromisoformat(t["due_date"]).date() >= today,
                datetime.fromisoformat(t["due_date"]).date()
            )
        )

        for index, task in enumerate(sorted_tasks):
            col1, col2 = st.columns([0.55, 3])

            with col1:
                # Color
                color_index = user_data["courses_list"].index(task["course"])
                color = user_data["courses_colors"][color_index]
                st.markdown(
                    f"""
                    <span style="
                        background-color: {color};
                        color: white;
                        padding: 5px 10px;
                        border-radius: 5px;
                        font-weight: bold;
                    ">
                        {task['course']}
                    </span>
                    """,
                    unsafe_allow_html=True,
                )

            with col2:

               # Unique key for each task
                expander_key = f"task_open_{index}"

                # Toggle button
                if expander_key not in st.session_state:
                    st.session_state[expander_key] = False

                if st.button(task["name"], key=f"toggle_{index}", use_container_width=True):
                    st.session_state[expander_key] = not st.session_state[expander_key]

                # Container for task content
                if st.session_state[expander_key]:
                    st.markdown('<div class="task-container">', unsafe_allow_html=True)

                    # ---- Task details go inside here ----
                    col1, col2, col3 = st.columns(3)

                    with col1:
                        task["name"] = st.text_input("Name", value=task["name"], key=f"name_{index}")
                        task["course"] = st.selectbox(
                            "Course",
                            user_data["courses_list"],
                            index=user_data["courses_list"].index(task["course"]),
                            key=f"course_{index}",
                        )
                        task["due_date"] = st.date_input("Due Date", value=task["due_date"], key=f"date_{index}").isoformat()

                    with col2:
                        task["status"] = st.select_slider(
                            "Status",
                            ["Not Started", "In-Progress", "Near Completion", "Complete"],
                            value=task["status"],
                            key=f"status_{index}",
                        )
                        task["effort"] = st.slider("Effort", min_value=1, max_value=5, value=task["effort"], key=f"effort_{index}")


                    

                    # --- AI or Manual Priority System (All Recalculate, Each Displays Its Own) ---

                    priority_prefix = "ai_priority_"
                    stale_key = "ai_data_stale_priority_all"

                    # --- Handle AI-based Priority Mode ---
                    if user_data.get("ai_use_ai_priority", False):

                        # Button (shown per task) — triggers all regeneration
                        if st.button("Regenerate Task Priority", key=f"regenerate_priority_btn_{index}"):
                            st.session_state[stale_key] = True

                        # Initialize state defaults
                        if stale_key not in st.session_state:
                            st.session_state[stale_key] = True

                        # --- Recalculate All Task Priorities if Needed ---
                        if st.session_state[stale_key]:
                            genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
                            model = genai.GenerativeModel(
                                model_name="gemini-2.5-flash",
                                generation_config=genai.GenerationConfig(temperature=0.0)
                            )

                            # Summarize all tasks for context
                            task_descriptions = []
                            for i, t in enumerate(user_data["tasks"]):
                                desc = f"{i+1}. {t['name']} (Course: {t['course']}, Difficulty: {t.get('difficulty', 'N/A')}, Due: {t['due_date']})"
                                task_descriptions.append(desc)
                            all_tasks_text = "\n".join(task_descriptions)

                            # AI prompt (batch mode)
                            prompt = f"""
                            You are an academic task prioritization assistant.

                            Your job is to assign a single-word **priority rating** ("High", "Medium", or "Low") to each of the following tasks.

                            You must ONLY output one word per task, separated by line breaks — in the same order the tasks are listed below.
                            Example Output:
                            High
                            Medium
                            Low
                            Medium

                            Rules:
                            1. **Due Date**
                            - Tasks due soon (today or within 3 days) → more likely "High"
                            - Tasks due in 4–7 days → usually "Medium"
                            - Tasks due in more than 7 days → likely "Low"

                            2. **Course Difficulty**
                            - Harder courses increase the likelihood of "High" or "Medium".
                            - Easier courses lower the likelihood of "High".

                            3. **Task Name**
                            - If the name includes words like "exam", "test", "quiz", "project", or "assignment", raise the priority.
                            - If the name includes words like "notes", "practice", "reading", or "review", lower the priority.

                            4. **Distribution Rule**
                            - On average, only 2 out of every 4 tasks (≈ 50%) should be rated "High".
                            - The rest must be "Medium" or "Low" according to importance.

                            Today's Date: {today}

                            Tasks:
                            {all_tasks_text}

                            Output one line per task (High / Medium / Low only):
                            """

                            with st.spinner("Regenerating priorities for all tasks..."):
                                try:
                                    response = model.generate_content(prompt)
                                    ai_output_lines = [
                                        line.strip() for line in response.text.splitlines()
                                        if line.strip() in ["High", "Medium", "Low"]
                                    ]

                                    for i, t in enumerate(user_data["tasks"]):
                                        # Assign the generated priority
                                        if i < len(ai_output_lines):
                                            t["priority"] = ai_output_lines[i]
                                        else:
                                            t["priority"] = "Medium"
                                        st.session_state[f"{priority_prefix}{i}"] = t["priority"]

                                except Exception as e:
                                    st.warning(f"AI failed to determine priorities: {e}")
                                    for i, t in enumerate(user_data["tasks"]):
                                        t["priority"] = "Medium"
                                        st.session_state[f"{priority_prefix}{i}"] = "Medium"

                            # Mark as no longer stale
                            st.session_state[stale_key] = False

                        # --- Display ONLY This Task’s Priority ---
                        task["priority"] = st.session_state.get(
                            f"{priority_prefix}{index}",
                            task.get("priority", "Medium")
                        )

                        color_priority = {
                            "Low": "green",
                            "Medium": "orange",
                            "High": "red",
                        }.get(task["priority"], "blue")

                        st.markdown(
                            f"""
                            <span style="
                                background-color: {color_priority};
                                color: white;
                                padding: 5px 10px;
                                border-radius: 5px;
                                font-weight: bold;
                            ">
                                {task['priority']}
                            </span>
                            """,
                            unsafe_allow_html=True,
                        )

                    else:
                        with col2:
                            # --- Manual Mode (Independent per task) ---
                            task["priority"] = st.select_slider(
                                "Priority",
                                ["Low", "Medium", "High"],
                                value=task.get("priority", "Medium"),
                                key=f"priority_slider_{index}",
                            )

                            color_priority = {
                                "Low": "green",
                                "Medium": "orange",
                                "High": "red",
                            }.get(task["priority"], "blue")

                            st.markdown(
                                f"""
                                <span style="
                                    background-color: {color_priority};
                                    color: white;
                                    padding: 5px 10px;
                                    border-radius: 5px;
                                    font-weight: bold;
                                ">
                                    {task['priority']}
                                </span>
                                """,
                                unsafe_allow_html=True,
                            )

                    # Notes section
                    task["written_notes"] = st.text_area(
                        "Notes",
                        value=task.get("written_notes", ""),
                        key=f"notes_{index}", height= 210
                    )

                    
                    
                    # File Upload
                
                    # Define the quiz dialog
                    def show_quiz_dialog(ai_output, task_name):
                        @st.dialog(f"{task_name} Quiz", width="large")
                        def view_quiz():
                            st.write(ai_output)
                        view_quiz()

                    def show_ai_summary(ai_output_summary):
                        @st.dialog(f"{task['uploaded_file_name']} Summary", width="large")
                        def view_ai_summary():
                            task["uploaded_file_name"] = uploaded_file.name
                            st.write(ai_output_summary)
                        view_ai_summary()


                    # Define the uploaded file preview dialog
                    def show_uploaded_file_dialog(uploaded_file):
                        @st.dialog("Uploaded File Preview", width="large")
                        def view_uploaded_file():
                            docx_bytes = uploaded_file.getvalue()
                            doc_stream = io.BytesIO(docx_bytes)
                            document = Document(doc_stream)

                            task["uploaded_file_name"] = uploaded_file.name

                            st.subheader("Document Content:")
                            for paragraph in document.paragraphs:
                                st.write(paragraph.text)

                            if document.tables:
                                st.subheader("Tables:")
                                for table in document.tables:
                                    for row in table.rows:
                                        row_data = [cell.text for cell in row.cells]
                                        st.write(row_data)


                        view_uploaded_file()

                                            

                    
                    with st.expander("Upload Notes"):
                        col9, col10 = st.columns(2)
                        with col9:
                            uploaded_file = st.file_uploader("Upload your file", type=["docx", "txt", "pdf"], key=f"uploaded_notes_{index}")


                            # Extract text content from the file
                            def extract_text(file):
                                if file.name.endswith(".docx"):
                                    doc = Document(file)
                                    return "\n".join([p.text for p in doc.paragraphs])
                                elif file.name.endswith(".txt"):
                                    return file.read().decode("utf-8")
                                elif file.name.endswith(".pdf"):
                                    # Optional: Add PDF extraction support using PyMuPDF or pdfminer
                                    return "PDF summarization not yet implemented."
                                else:
                                    return "Unsupported file type."

                            # Handle file upload
                            if uploaded_file:
                                content = extract_text(uploaded_file)
                            else:
                                content = "No Uploaded File"

                        
                        if uploaded_file:

                        

                            with col10:
                                task['uploaded_file_name'] = uploaded_file.name
                                st.info(f"Uploaded: {task['uploaded_file_name']}")

                                if st.button(f"Preview {task['uploaded_file_name']}", key=f"preview_{index}"):
                                    show_uploaded_file_dialog(uploaded_file)

                                
                            st.header("AI Tools")
                        
                            def generate_quiz_button():
                                if st.button("Generate Quiz", key=f"generate_quiz_button_{index}"):

                                    # Configure Gemini
                                    genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
                                    model = genai.GenerativeModel(
                                        model_name="gemini-2.5-flash",
                                        generation_config=genai.GenerationConfig(temperature=1.2)
                                    )

                                    prompt = f"""
                                    Your task is to output a quiz for the user based off of the notes they submit

                                    Take the name of the task, the course, the users difficulty ranking of the course and the rest of the task data into consideration.
                                    You will take all that info as well as the file uploaded as notes and generate a quiz relevant to what is in the notes, as well as some questions from other internet souces relevant to the topic
                                    Use the AI Input Guidelines to determine the type and length of the quiz.


                                    INPUTS
                                    difficulty: {difficulty_ranking}
                                    courses: {courses}
                                    tasks: {tasks}
                                    file: {content}

                                    AI Input Guidelines:
                                    Response Length: {ai_quiz_length}
                                    Short: - 1 to 5 questions (Multiple Choice)    
                                    Medium: - 1 to 10 questions  (2/3 Multiple Choice, 1/3 Short Answer)                 
                                    Long: - 1 to 20 questions  (2/3 Multiple Choice, 1/3 Short Answer)     
                                    Complete Review: - 1 to 25 questions (1/2 Multiple Choice, 1/2 Short Answer)         
                                    


                                    Formating requirements:
                                    Generate a 20 or so question quiz on the uploaded notes (uploaded file) and other KNOWN AND CREDIBLE SOURCES
                                    If relevant, create mulitple choice questions where the options are easily seen in the question and word questions for relevant topics
                                    Order them in easiest to hardest as the quiz goes along

                                    ONLY PROVIDE QUETSIONS THAT HAVE ACTUAL, REAL ANSWERS
                                    """

                                    with st.spinner("Generating Quiz ..."):
                                        response = model.generate_content(prompt)
                                        ai_output = response.text
                                        st.session_state["ai_quiz"] = [ai_output]
                                        st.success("Quiz Successfully Generated")

                                        # Trigger the quiz dialog
                                        show_quiz_dialog(ai_output, task["name"])

                        
                            def summary_button():
                                if st.button("Summarize", key=f"summarize_button_{index}"):
                                    # Configure Gemini
                                    genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
                                    model = genai.GenerativeModel(
                                        model_name="gemini-2.5-flash",
                                        generation_config=genai.GenerationConfig(temperature=1.0)
                                    )

                                    prompt = f"""

                                    You are an AI assistant designed to read and summarize documents submitted by users. Your goal is to generate a clear, concise, and informative summary of the uploaded file. Keep the summary brief, but do not omit important details or critical context.
                                    
                                    Instructions:
                                    - Read and analyze the document thoroughly.
                                    - Extract the main topics, arguments, conclusions, and any key data points or examples.
                                    - Prioritize clarity and brevity while maintaining accuracy and completeness.
                                    - Do not copy large chunks of text from the document.
                                    - Do not add external information not found in the file.
                                    - Use the AI Input Guidelines to determine how long and thus detailed the summary is.
                                    
                                    Output Format:
                                    - Title (if available)
                                    - Brief overview (1-5 sentences)

                                    Use AI Input Guidelines at this point to determine how long and how much detail to include. Still use the following 3 guildlines as a maximum output gridline.
                                    - Bullet-point summary of key points (5-12 items max)
                                    - Include sections/quotation of the document where the summary points were taken from
                                    - [Optional] Definitions or explanations of complex terms, if needed for clarity

                                    Example Output (for a class note or report):
                                    Title: Introduction to Quantum Mechanics
                                    Overview: This document provides a basic introduction to the fundamental concepts of quantum mechanics.
                                    Key Points:
                                    Quantum mechanics describes the behavior of particles at atomic and subatomic scales.
                                    Wave-particle duality means particles can behave like waves and vice versa.
                                    The Schrödinger equation is used to predict how quantum systems evolve.
                                    Measurement affects the system being observed (observer effect).
                                    Applications include quantum computing, cryptography, and semiconductors.
                                    
                                    Inputs:
                                    difficulty: {difficulty_ranking}
                                    courses: {courses}
                                    tasks: {tasks}
                                    file: {content}

                                    AI Input Guidelines:
                                    Response length: {ai_summary_length}
                                    Begin your summary below:
                                    

                                    """
                                    with st.spinner("Generating Summary ..."):
                                        response = model.generate_content(prompt)
                                        ai_output_summary = response.text
                                        st.session_state["ai_summary"] = [ai_output_summary]
                                        st.success("Summary Successfully Generated")

                                        # Trigger the quiz dialog
                                        show_ai_summary(ai_output_summary)

                            if st.session_state["current_ai_session"]:
                                col6, col7, col8= st.columns(3)

                                with col6:
                                    generate_quiz_button()

                                with col7:
                                    summary_button()

                                with col8:
                                    if st.button("Clear Chat Session"):
                                        st.session_state["current_ai_session"] = []
                            else:
                                col6, col7,= st.columns(2)

                                with col6:
                                    generate_quiz_button()

                                with col7:
                                    summary_button()



                    if user_data['ai_document_assistant']:        
                        # Configure Gemini API
                        genai.configure(api_key=st.secrets["GEMINI_API_KEY"])

                        # Initialize Gemini Model 
                        model = genai.GenerativeModel(
                            model_name="gemini-2.5-flash",
                            generation_config=genai.GenerationConfig(temperature=1.1)
                        )

                        # Load AI History Only If Enabled
                        if user_data.get("ai_use_history", False):
                            for entry in st.session_state.get("current_ai_session", []):
                                with st.chat_message("user"):
                                    st.markdown(entry["user_input"])
                                with st.chat_message("assistant"):
                                    st.markdown(entry["ai_response"])

                        # AI Chat Input 
                        user_input = st.chat_input("Prompt AI Assistant")

                        # Conditionally include AI history in the prompt 
                        the_ai_history = (
                            user_data.get("ai_history", []) if user_data.get("ai_use_history", False) else []
                        )

                        # Ai Prompt
                        prompt = f"""
                        You are a friendly, supportive, and knowledgeable AI study assistant helping students learn effectively.

                        Speak in an approachable, encouraging tone — like a helpful tutor or study buddy. Be patient, clear, and non-judgmental.

                        Your main goals:
                        - Explain complex topics in simple, easy-to-understand language
                        - Provide summaries, examples, and memory aids (like mnemonics)
                        - Use the AI Input Guideline to determine the length and detail of your responses
                        - Ask clarifying questions when a student's request is unclear
                        - Help students build study skills, time management, and focus
                        - Support motivation and celebrate effort, not just correct answers
                        - Use analogies, real-world examples, or relatable explanations to engage students

                        Avoid overly formal or robotic language. Always prioritize learning, understanding, and encouragement.

                        Use the AI chat history for context (if applicable), as well as the course list and difficulty settings to adapt explanations.
                        You can also answer questions about the attached document.
                        And of course, respond helpfully to the user prompt.

                        Inputs:
                        difficulty: {difficulty_ranking}
                        courses: {courses}
                        tasks: {tasks}
                        ai chat history: {the_ai_history}
                        attached document: {content}
                        user input prompt: {user_input}

                        AI Input Guidelines:
                        Response length: {ai_assistant_response_length}
                        """

                        # Process User Input 
                        if user_input:
                            with st.chat_message("user"):
                                st.markdown(user_input)

                            with st.spinner("Generating Response"):
                                response = model.generate_content(prompt)
                                ai_response = response.text

                            with st.chat_message("assistant"):
                                st.markdown(ai_response)

                            # Save AI History Only If Enabled 
                            if user_data.get("ai_use_history", False):
                                new_entry = {
                                    "user_input": user_input,
                                    "ai_response": ai_response,
                                    "course": task['course'] if 'task' in locals() else None,
                                    "timestamp": datetime.utcnow().isoformat()
                                }

                                st.session_state.setdefault("current_ai_session", []).append(new_entry)
                                user_data.setdefault("ai_history", []).append(new_entry)
                                update_user_data(email, user_data)

                      

                            

                        

                        


                    # --- Save Updates ---
                    with col3:
                        # Update Task Button
                        if st.button(f"Update Task {index + 1}", key=f"update_button_{index}"):
                            st.session_state[f"ai_priority_stale_{index}"] = True
                            # user_data["tasks"][index] = task
                            update_user_data(email, user_data)
                            st.session_state["ai_data_stale"] = True
                            st.session_state["ai_data_stale_priority"] = True
                            

                        # Mark Complete
                        if st.button("Mark As Complete", key=f"mark_complete_button_{index}"):
                            # Make sure we have a place to store completed tasks
                            user_data.setdefault("complete_tasks", [])

                            # Move this specific task to completed
                            completed_task = user_data["tasks"].pop(index)
                            completed_task["status"] = "Complete"
                            user_data["complete_tasks"].append(completed_task)

                            # Update database
                            update_user_data(email, user_data)

                            # Mark AI data as stale so priorities refresh properly
                            st.session_state["ai_data_stale"] = True
                            st.session_state["ai_data_stale_priority"] = True

                            # Force page refresh after the move to avoid iterator conflicts

                        # Delete Task
                        if st.button(f"Delete Task {index + 1}", key=f"delete_{index}"):
                            if "uploaded_file_path" in task and os.path.exists(task["uploaded_file_path"]):
                                os.remove(task["uploaded_file_path"])

                            del user_data["tasks"][index]
                            update_user_data(email, user_data)
                            st.session_state["ai_data_stale"] = True
                            st.session_state["ai_data_stale_priority"] = True
                    
                    # Close content div
                    st.markdown("</div>", unsafe_allow_html=True)

    def ai_to_do_list():
    
        # Set default session variables
        # if "ai_data_stale" not in st.session_state:
        #     st.session_state["ai_data_stale"] = True

        if "ai_to_do_list_database" not in st.session_state:
            st.session_state["ai_to_do_list_database"] = []
        
        if "current_ai_session" not in st.session_state:
            st.session_state["current_ai_session"] = []


        # Configure the Gemini API
        genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
        model = genai.GenerativeModel(
            model_name="gemini-2.5-flash",
            generation_config=genai.GenerationConfig(temperature=0.0)
        )

        prompt = f"""
        Your task is to output a list of which of the users tasks in the logical order that they should be done in the most time efficient manor.

        Take the courses list and course difficulty as well as the tasks and their details in the database and create a list of which order to do the tasks in that is personalized to each users details.
        Take the names as well as the status of the tasks into consideration as well to determine the gravity of the task when determining the order.

        INPUTS
        difficulty: {difficulty_ranking}
        courses: {courses}
        tasks: {tasks}

        Formating requirements:
        Make it into a list that can be displayed in streamlit application
        YOU ARE ONLY OUTPUTING A LIST OF THE USER'S TASKS IN THE ORDER OF WHICH THEY SHOULD BE DONE BASED ON THE INPUTS
        DO NOT OUTPUT CODE OF ANY KIND, JUST THE LIST OF THE NAMES OF THE TASKS!!!!
        Do Not add extra unecessary characters like brackets in response, but ensure it remains a list, make it as easy as possible to incorporate response into a list

        What You Need To Output:
        - OUTPUT A POINT LIST THAT LISTS THE ORDER OF THE TASKS
        - Put them in a "To Do Today" category "To Do Afterwards" category and add a bit of space in between the two categories for easy legibility
        - If no other tasks after "To Do Today" category, do NOT include the "To Do Afterwards" category
        - If No current tasks at all in the database, just output "No Tasks To Order"
        - ONLY If tasks have the exact same name, stick the name of their respective course in brackets when listing them in the completion order
        - OUTPUT AS QUICKLY AS POSSIBLE WHILE STILL BEING AS ACCURATE AS POSSIBLE
        """

        # If data was updated and AI list is stale, regenerate
        if st.session_state["ai_data_stale"]:
            with st.spinner("Generating Optimal Task Completion Order ..."):
                response = model.generate_content(prompt)
                ai_output = response.text
                st.session_state["ai_to_do_list_database"] = [ai_output]
                st.session_state["ai_data_stale"] = False  # Mark as up-to-date
                st.success("AI To-Do List Updated!")
        
        # Show current AI-generated list
        if st.session_state["ai_to_do_list_database"]:
            st.markdown(st.session_state["ai_to_do_list_database"][0])

        

        # Manual update option
        if st.button("Regenerate AI Task Order", key = "generate_ai_task_order_button"):
            st.session_state["ai_data_stale"] = True

    if not courses:
        st.warning("Please Enter Your Courses In The Setting Menu Before Continuing")


 # Show the current tasks with AI to do list next to it if enabled
    
        

    with st.container(border=True):
            # Display The AI Ordered Task Completion List
            st.subheader("AI Workflow Optimizer")
            ai_to_do_list()


    # Display the list of tasks for editing
    st.subheader("Current Tasks")
    if user_data['tasks']:
        
        display_tasks()
    else:
        st.write("No Tasks Available.")


        






ai(st.session_state['username'])